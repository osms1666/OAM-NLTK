"""Document ingestion — PDF (with OCR fallback), DOCX, TXT, HTML."""
from __future__ import annotations

import logging
import re
from dataclasses import dataclass, field
from datetime import datetime
from pathlib import Path
from typing import Iterable

import pandas as pd

log = logging.getLogger(__name__)

SUPPORTED = {".pdf", ".docx", ".txt", ".md", ".html", ".htm"}


@dataclass
class Document:
    path: str
    name: str
    text: str
    language: str | None = None
    n_chars: int = 0
    n_tokens: int = 0
    date: str | None = None        # ISO date if parseable from filename/metadata
    country: str | None = None     # extracted from filename when possible
    meta: dict = field(default_factory=dict)

    def to_dict(self) -> dict:
        return {
            "path": self.path, "name": self.name, "text": self.text,
            "language": self.language, "n_chars": self.n_chars,
            "n_tokens": self.n_tokens, "date": self.date, "country": self.country,
            **self.meta,
        }


# ---------- low-level readers ------------------------------------------------

def _read_txt(path: Path) -> str:
    return path.read_text(encoding="utf-8", errors="ignore")


def _read_html(path: Path) -> str:
    from bs4 import BeautifulSoup
    html = path.read_text(encoding="utf-8", errors="ignore")
    return BeautifulSoup(html, "html.parser").get_text(" ", strip=True)


def _read_docx(path: Path) -> str:
    import docx  # python-docx
    d = docx.Document(str(path))
    return "\n".join(p.text for p in d.paragraphs)


def _read_pdf(path: Path, ocr: bool = True) -> str:
    """Extract text from a PDF. Falls back to OCR if the PDF is image-based."""
    try:
        from pdfminer.high_level import extract_text
        text = extract_text(str(path)) or ""
    except Exception as e:  # pragma: no cover
        log.warning("pdfminer failed on %s: %s", path, e)
        text = ""
    if len(text.strip()) > 50 or not ocr:
        return text
    # OCR fallback
    try:
        from pdf2image import convert_from_path
        import pytesseract
        pages = convert_from_path(str(path), dpi=300)
        return "\n".join(pytesseract.image_to_string(p, lang="eng+fra+ara") for p in pages)
    except Exception as e:  # pragma: no cover
        log.warning("OCR failed on %s: %s", path, e)
        return text


_READERS = {
    ".txt": _read_txt, ".md": _read_txt,
    ".html": _read_html, ".htm": _read_html,
    ".docx": _read_docx,
    ".pdf": _read_pdf,
}


# ---------- public API --------------------------------------------------------

_DATE_RE = re.compile(r"(19|20)\d{2}(?:[-_](0[1-9]|1[0-2]))?")


def _guess_date(name: str) -> str | None:
    m = _DATE_RE.search(name)
    if not m:
        return None
    try:
        return datetime.strptime(m.group(0).replace("_", "-"), "%Y-%m").date().isoformat()
    except ValueError:
        try:
            return datetime.strptime(m.group(0)[:4], "%Y").date().isoformat()
        except ValueError:
            return None


def _guess_country(name: str) -> str | None:
    # "NDC_Kenya.pdf" → "Kenya"
    m = re.search(r"[_\- ]([A-Z][a-zA-Z]{2,})(?:\.|_|$)", name)
    return m.group(1) if m else None


def load_document(path: str | Path, ocr: bool = True) -> Document:
    p = Path(path)
    ext = p.suffix.lower()
    if ext not in _READERS:
        raise ValueError(f"Unsupported file type: {ext}")
    text = _READERS[ext](p) if ext != ".pdf" else _read_pdf(p, ocr=ocr)
    doc = Document(
        path=str(p.resolve()),
        name=p.stem,
        text=text or "",
        n_chars=len(text or ""),
        n_tokens=len((text or "").split()),
        date=_guess_date(p.name),
        country=_guess_country(p.name),
    )
    try:
        from langdetect import detect
        if doc.text.strip():
            doc.language = detect(doc.text[:2000])
    except Exception:
        pass
    return doc


def load_folder(
    folder: str | Path,
    recursive: bool = True,
    ocr: bool = True,
    patterns: Iterable[str] | None = None,
) -> pd.DataFrame:
    """Load every supported document in a folder into a DataFrame."""
    root = Path(folder)
    if not root.exists():
        raise FileNotFoundError(root)
    glob = "**/*" if recursive else "*"
    files = [p for p in root.glob(glob) if p.is_file() and p.suffix.lower() in SUPPORTED]
    if patterns:
        files = [p for p in files if any(p.match(pat) for pat in patterns)]
    docs = []
    for p in files:
        try:
            docs.append(load_document(p, ocr=ocr).to_dict())
        except Exception as e:  # pragma: no cover
            log.warning("Failed %s: %s", p, e)
    return pd.DataFrame(docs)
